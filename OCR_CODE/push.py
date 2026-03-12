import streamlit as st
import os
import re
import json
import requests
import tempfile
from datetime import datetime
from docx import Document
from typing import Dict, List, Any, Optional
import io

class StudentProfileExtractor:
    """Extract student profile information from document"""
    def __init__(self, doc_content):
        self.doc = Document(io.BytesIO(doc_content))
        self.full_text = self.extract_text_from_doc()
        
    def extract_text_from_doc(self) -> str:
        """Extract all text from the document"""
        full_text = []
        for paragraph in self.doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in self.doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return '\n'.join(full_text)
    
    def extract_pincode_from_address(self, address: str) -> str:
        """Extract 6-digit pincode from address"""
        pincode_matches = re.findall(r'\b(\d{6})\b', address)
        return pincode_matches[0] if pincode_matches else "400001"
    
    def extract_city_state_from_address(self, address: str) -> tuple:
        """Extract city and state from address"""
        address_lower = address.lower()
        
        city_state_map = {
            'delhi': ('Delhi', 'Delhi'), 
            'new delhi': ('New Delhi', 'Delhi'),
            'noida': ('Noida', 'Uttar Pradesh'), 
            'gurgaon': ('Gurgaon', 'Haryana'), 
            'gurugram': ('Gurugram', 'Haryana'), 
            'faridabad': ('Faridabad', 'Haryana'),
            'ghaziabad': ('Ghaziabad', 'Uttar Pradesh'), 
            'mumbai': ('Mumbai', 'Maharashtra'),
            'pune': ('Pune', 'Maharashtra'), 
            'bangalore': ('Bangalore', 'Karnataka'),
            'bengaluru': ('Bengaluru', 'Karnataka'), 
            'hyderabad': ('Hyderabad', 'Telangana'),
            'chennai': ('Chennai', 'Tamil Nadu'), 
            'kolkata': ('Kolkata', 'West Bengal'),
            'ahmedabad': ('Ahmedabad', 'Gujarat'), 
            'jaipur': ('Jaipur', 'Rajasthan'),
            'lucknow': ('Lucknow', 'Uttar Pradesh'), 
            'kanpur': ('Kanpur', 'Uttar Pradesh'),
            'nagpur': ('Nagpur', 'Maharashtra'), 
            'indore': ('Indore', 'Madhya Pradesh'),
            'bhopal': ('Bhopal', 'Madhya Pradesh'), 
            'patna': ('Patna', 'Bihar'),
            'vadodara': ('Vadodara', 'Gujarat'), 
            'ludhiana': ('Ludhiana', 'Punjab')
        }
        
        for city_name, (city, state) in city_state_map.items():
            if city_name in address_lower:
                return city, state
        
        return "Mumbai", "Maharashtra"  # Default
    
    def parse_date(self, date_str: str) -> str:
        """Parse date string to YYYY-MM-DD format"""
        if not date_str:
            return None
            
        # Clean the date string
        clean_date = ''.join(c for c in date_str if c.isdigit() or c in '/-.')
        
        # Try different date patterns
        date_patterns = [
            r'(\d{1,2})[/-](\d{1,2})[/-](\d{4})',  # DD/MM/YYYY
            r'(\d{1,2})[/-](\d{1,2})[/-](\d{2})',   # DD/MM/YY
            r'(\d{4})[/-](\d{1,2})[/-](\d{1,2})',   # YYYY/MM/DD
        ]
        
        for pattern in date_patterns:
            match = re.search(pattern, clean_date)
            if match:
                part1, part2, part3 = match.groups()
                
                if len(part1) == 4:  # YYYY format
                    year, month, day = part1, part2, part3
                else:  # DD/MM format
                    day, month, year = part1, part2, part3
                    if len(year) == 2:
                        year = "20" + year if int(year) <= 50 else "19" + year
                
                return f"{year}-{month.zfill(2)}-{day.zfill(2)}"
        
        return None
    
    def extract_student_profile(self) -> Dict[str, Any]:
        """Extract student profile information from ASSESSMENT FORM"""
        lines = self.full_text.split('\n')
        
        # Initialize variables
        student_name = "Unknown Student"
        dob = None
        doa = None
        gender = "M"
        mother_tongue = "N/A"
        languages_spoken = "N/A"
        address = ""
        condition = ""
        phone_residential = "+91-1234567890"  # Default
        emergency_contact_name = "Emergency Contact"
        emergency_phone = "+91-1234567890"  # Default
        
        # First, try to find ASSESSMENT FORM or ADMISSION FORM section
        in_assessment_form = False
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            
            # Look for ASSESSMENT FORM or ADMISSION FORM
            if any(form in line_stripped.upper() for form in ['ASSESSMENT FORM', 'ADMISSION FORM']):
                in_assessment_form = True
                continue
            
            if in_assessment_form:
                # Extract Name - on its own line like "Name: Anrar Kakkar"
                if line_stripped.startswith('Name:') and student_name == "Unknown Student":
                    name_value = line_stripped.replace('Name:', '').strip()
                    if name_value:
                        student_name = name_value
                
                # Extract D.O.B - on its own line like "D.O.B.: 27-03-2012"
                elif line_stripped.startswith('D.O.B') and not dob:
                    dob_match = re.search(r'D\.O\.B\.?:?\s*([0-9/-]+)', line_stripped)
                    if dob_match:
                        dob_str = dob_match.group(1).strip()
                        dob = self.parse_date(dob_str)
                
                # Extract Sex/Gender - on its own line like "Sex: Male"
                elif line_stripped.startswith('Sex:') or line_stripped.startswith('Gender:'):
                    sex_match = re.search(r'(?:Sex|Gender):\s*(\w+)', line_stripped)
                    if sex_match:
                        gender_value = sex_match.group(1).strip().upper()
                        if gender_value.startswith('M'):
                            gender = "M"
                        elif gender_value.startswith('F'):
                            gender = "F"
                
                # Extract D.O.A - on its own line like "D.O.A.: 20-07-2022"
                elif line_stripped.startswith('D.O.A') and not doa:
                    doa_match = re.search(r'D\.O\.A\.?:?\s*([0-9/-]+)', line_stripped)
                    if doa_match:
                        doa_str = doa_match.group(1).strip()
                        doa = self.parse_date(doa_str)
                
                # Extract Mother Tongue - on its own line like "Mother Tongue: Hindi"
                elif line_stripped.startswith('Mother Tongue:'):
                    tongue_value = line_stripped.replace('Mother Tongue:', '').strip()
                    if tongue_value:
                        mother_tongue = tongue_value
                
                # Extract Languages Spoken
                elif line_stripped.startswith('Languages Spoken:'):
                    lang_value = line_stripped.replace('Languages Spoken:', '').strip()
                    if lang_value:
                        languages_spoken = lang_value
                
                # Extract Residential Address - on its own line
                elif line_stripped.startswith('Residential Address:'):
                    address = line_stripped.replace('Residential Address:', '').strip()
                
                # Extract Phone No. (could be residential) - on its own line like "Phone No.: 7982269508"
                elif line_stripped.startswith('Phone No.:') and phone_residential == "+91-1234567890":
                    phone_match = re.search(r'Phone No\.?:?\s*([0-9\s/-]+)', line_stripped)
                    if phone_match:
                        phone_digits = re.sub(r'[^0-9]', '', phone_match.group(1))
                        if len(phone_digits) == 10:
                            phone_residential = f"+91-{phone_digits}"
                
                # Extract Condition - on its own line like "Condition: Moderate Autism"
                elif line_stripped.startswith('Condition:'):
                    condition = line_stripped.replace('Condition:', '').strip()
                
                # Extract Alternative Contact Name (emergency contact)
                elif line_stripped.startswith('Alternative Contact Name:'):
                    emergency_contact_name = line_stripped.replace('Alternative Contact Name:', '').strip()
                
                # Extract Tel. (emergency phone) - on its own line like "Tel.: 9810284465 / 9818418511"
                elif line_stripped.startswith('Tel.:'):
                    tel_match = re.search(r'Tel\.?:?\s*([0-9\s/-]+)', line_stripped)
                    if tel_match:
                        tel_value = tel_match.group(1).strip()
                        phone_digits = re.sub(r'[^0-9]', '', tel_value.split('/')[0])
                        if len(phone_digits) == 10:
                            emergency_phone = f"+91-{phone_digits}"
                
                # Stop at next major section
                if any(section in line_stripped.upper() for section in [
                    'FAMILY CONSTELLATION', 
                    'CHIEF COMPLAINTS', 
                    'PROBLEM IN DETAIL',
                    'PROBLEMS IN DETAIL',
                    'ON EXAMINATION',
                    'CHIEF ASSESSMENT',
                    'OT/PT ASSESSMENT'
                ]):
                    break
        
        # If we didn't find in assessment form, search entire document
        if student_name == "Unknown Student":
            for line in lines:
                # Try to find Name anywhere
                if line.strip().startswith('Name:'):
                    name_value = line.replace('Name:', '').strip()
                    if name_value:
                        student_name = name_value
                        break
        
        if not dob:
            for line in lines:
                if 'D.O.B' in line or 'DOB' in line.upper():
                    dob_match = re.search(r'D\.?O\.?B\.?:?\s*([0-9/-]+)', line, re.IGNORECASE)
                    if dob_match:
                        dob_str = dob_match.group(1).strip()
                        dob = self.parse_date(dob_str)
                        break
        
        if not doa:
            for line in lines:
                if 'D.O.A' in line or 'DOA' in line.upper():
                    doa_match = re.search(r'D\.?O\.?A\.?:?\s*([0-9/-]+)', line, re.IGNORECASE)
                    if doa_match:
                        doa_str = doa_match.group(1).strip()
                        doa = self.parse_date(doa_str)
                        break
        
        if not condition:
            for line in lines:
                if line.strip().startswith('Condition:'):
                    condition = line.replace('Condition:', '').strip()
                    break
        
        if not mother_tongue or mother_tongue == "N/A":
            for line in lines:
                if line.strip().startswith('Mother Tongue:'):
                    mother_tongue = line.replace('Mother Tongue:', '').strip()
                    break
        
        # Parse student name into first and last name
        name_parts = student_name.strip().split()
        if len(name_parts) == 0:
            first_name = "Unknown"
            last_name = "Student"
        elif len(name_parts) == 1:
            first_name = name_parts[0]
            last_name = "Student"
        elif len(name_parts) == 2:
            first_name = name_parts[0]
            last_name = name_parts[1]
        else:
            # For names with 3+ parts (e.g., "Atul Kumar Shah")
            first_name = name_parts[0]
            last_name = name_parts[-1]
        
        # If DOA not found, use current date
        if not doa:
            doa = datetime.now().strftime("%Y-%m-%d")
        
        # Extract city, state, pincode from address
        city, state = self.extract_city_state_from_address(address) if address else ("Mumbai", "Maharashtra")
        pin_code = self.extract_pincode_from_address(address) if address else "400001"
        
        # Create email from first name only
        email_name = first_name.lower().replace(" ", ".")
        email = f"{email_name}@gmail.com"
        
        student_data = {
            "first_name": first_name,
            "last_name": last_name,
            "date_of_birth": dob or "2013-07-28",
            "gender": gender,
            "date_of_admission": doa,
            "status": "Active",
            "address": address if address else "",
            "city": city,
            "state": state,
            "pin_code": pin_code,
            "phone": phone_residential,
            "email": email,
            "mother_tongue": mother_tongue,
            "languages_spoken": languages_spoken,
            "primary_condition": condition if condition else "Autism Spectrum Disorder",
            "emergency_contact": emergency_contact_name,
            "emergency_phone": emergency_phone,
            "program": "9c82dadb-5b2b-4341-ad91-498d5ca02fd5",
            "date_enrolled": doa,
            "uses_transport": "true",
            "photo": None
        }
        
        return student_data

class AssessmentDataExtractor:
    def __init__(self, doc_content):
        self.doc = Document(io.BytesIO(doc_content))
        self.extracted_data = {}
        self.full_text = self.extract_text_from_doc()
        self.lines = self.full_text.split('\n')
        
    def extract_text_from_doc(self) -> str:
        """Extract all text from the document"""
        full_text = []
        for paragraph in self.doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in self.doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return '\n'.join(full_text)
    
    def extract_value_from_line(self, line: str, field_name: str) -> str:
        """Extract value from a line containing field name"""
        if field_name.lower() in line.lower():
            parts = line.split(':', 1)
            if len(parts) > 1:
                return parts[1].strip()
        return ""
    
    def extract_section_content(self, section_name: str, max_lines: int = 20) -> str:
        """Extract content of a section until next section"""
        content = []
        in_section = False
        lines_collected = 0
        
        for line in self.lines:
            line_stripped = line.strip()
            
            # Check if this line starts the section
            if section_name.upper() in line_stripped.upper():
                in_section = True
                continue
            
            if in_section:
                # Check if we've reached another section
                if (line_stripped and 
                    (line_stripped.isupper() and len(line_stripped) < 100) or
                    any(sec in line_stripped.upper() for sec in [
                        'ASSESSMENT', 'FAMILY', 'CHIEF', 'PROBLEM', 'ON EXAMINATION',
                        'DEVELOPMENTAL', 'GEN EXAMINATION', 'MEDICAL', 'RECOMMENDATIONS',
                        'DIAGNOSTIC', 'PROVISIONAL', 'SIGNATURE', 'CONCLUSION'
                    ])):
                    break
                
                if line_stripped:
                    content.append(line_stripped)
                    lines_collected += 1
                    if lines_collected >= max_lines:
                        break
        
        return ' '.join(content) if content else ""
    
    def find_field_in_lines(self, field_names: List[str]) -> str:
        """Find field value in lines"""
        for line in self.lines:
            for field in field_names:
                if field.lower() in line.lower():
                    parts = line.split(':', 1)
                    if len(parts) > 1:
                        return parts[1].strip()
        return ""
    
    def normalize_quality(self, value: str) -> str:
        """Normalize quality values to good/fair/poor"""
        if not value:
            return ""
        
        value_lower = value.lower()
        
        if "good" in value_lower:
            return "good"
        elif "fair" in value_lower:
            return "fair"
        elif "poor" in value_lower:
            return "poor"
        elif "normal" in value_lower or "n" == value_lower:
            return "fair"
        elif "intact" in value_lower:
            return "fair"
        elif "poor-fair" in value_lower:
            return "fair"
        elif "fair-good" in value_lower:
            return "good"
        elif "normal-fair" in value_lower:
            return "fair"
        else:
            return value
    
    def parse_date_from_text(self, text: str) -> str:
        """Parse date from text"""
        date_patterns = [
            r'(\d{1,2})[/-](\d{1,2})[/-](\d{4})',
            r'(\d{1,2})[/-](\d{1,2})[/-](\d{2})',
            r'(\d{4})[/-](\d{1,2})[/-](\d{1,2})',
        ]
        
        for pattern in date_patterns:
            match = re.search(pattern, text)
            if match:
                part1, part2, part3 = match.groups()
                if len(part1) == 4:
                    year, month, day = part1, part2, part3
                else:
                    day, month, year = part1, part2, part3
                    if len(year) == 2:
                        year = "20" + year if int(year) <= 50 else "19" + year
                return f"{year}-{month.zfill(2)}-{day.zfill(2)}"
        return ""
    
    def extract_assessment_data(self) -> Dict[str, Any]:
        """Extract assessment data"""
        # Get assessment date from D.O.A
        assessment_date = ""
        for line in self.lines:
            if 'D.O.A' in line or 'DOA' in line.upper():
                date_match = re.search(r'D\.?O\.?A\.?:?\s*([0-9/-]+)', line, re.IGNORECASE)
                if date_match:
                    assessment_date = self.parse_date_from_text(date_match.group(1))
                    break
        
        if not assessment_date:
            assessment_date = datetime.now().strftime("%Y-%m-%d")
        
        # Extract conclusion
        conclusion = ""
        for i, line in enumerate(self.lines):
            if 'conclusion' in line.lower():
                # Get next few lines until signature or end
                conclusion_lines = []
                for j in range(i+1, min(i+10, len(self.lines))):
                    next_line = self.lines[j].strip()
                    if not next_line or 'signature' in next_line.lower():
                        break
                    conclusion_lines.append(next_line)
                conclusion = ' '.join(conclusion_lines)
                break
        
        # Extract strengths from special abilities
        strengths = ""
        for i, line in enumerate(self.lines):
            if 'special ability' in line.lower() or 'special abilities' in line.lower():
                # Get content after colon or next lines
                if ':' in line:
                    strength_part = line.split(':', 1)[1].strip()
                    if strength_part:
                        strengths = strength_part
                else:
                    # Check next few lines
                    for j in range(i+1, min(i+5, len(self.lines))):
                        next_line = self.lines[j].strip()
                        if not next_line or 'conclusion' in next_line.lower():
                            break
                        strengths += " " + next_line
                    strengths = strengths.strip()
                break
        
        # Extract recommendations from language development
        recommendations = ""
        for i, line in enumerate(self.lines):
            if 'recommendations' in line.lower():
                # Check if it's in language development context
                context_lines = self.lines[max(0, i-5):i]
                context_text = ' '.join(context_lines).lower()
                if 'language' in context_text or 'speech' in context_text:
                    rec_lines = []
                    for j in range(i+1, min(i+15, len(self.lines))):
                        next_line = self.lines[j].strip()
                        if not next_line or any(sec in next_line.lower() for sec in ['signature', 'diagnostic', 'provisional']):
                            break
                        rec_lines.append(next_line)
                    recommendations = ' '.join(rec_lines)
                    break
        
        # Extract challenges from Condition field
        challenges = ""
        for line in self.lines:
            if 'Condition:' in line:
                challenge_part = line.split(':', 1)[1].strip()
                if challenge_part:
                    challenges = challenge_part
                    # Clean up - remove any other field names
                    for field in ['Associated Condition', 'Waiting list', 'Call done']:
                        if field in challenges:
                            challenges = challenges.split(field)[0].strip()
                    break
        
        assessment_data = {
            "assessment_type": "Initial",
            "assessment_date": assessment_date,
            "location": "MBCN",
            "student": "",
            "team": "1863c310-743f-4f17-a630-62f17ba8fc1d",
            "waitlist_entry": None,
            "assessors": [1],
            "summary": conclusion if conclusion else "N/A",
            "strengths": strengths if strengths else "N/A",
            "challenges": challenges if challenges else "N/A",
            "recommendations": recommendations if recommendations else "Speech therapy: [No information available] Parent education: [No information available]",
            "decision": "Admit",
            "recommended_program": "",
            "decision_date": assessment_date,
            "follow_up_date": "2025-12-15",
            "is_completed": False
        }
        
        self.extracted_data['assessment'] = assessment_data
        return assessment_data
    
    def extract_pre_assessment_data(self) -> Dict[str, Any]:
        """Extract pre-assessment data"""
        problems_details = ""
        
        # Look for PROBLEMS IN DETAIL section
        for i, line in enumerate(self.lines):
            if 'PROBLEMS IN DETAIL:' in line.upper() or 'PROBLEM IN DETAIL:' in line.upper():
                # Extract content after the heading
                content_lines = []
                for j in range(i+1, min(i+30, len(self.lines))):
                    next_line = self.lines[j].strip()
                    if not next_line or any(sec in next_line.upper() for sec in [
                        'DEVELOPMENTAL', 'GEN EXAMINATION', 'MEDICAL', 'FAMILY',
                        'CHIEF COMPLAINTS', 'ON EXAMINATION'
                    ]):
                        break
                    content_lines.append(next_line)
                problems_details = ' '.join(content_lines)
                break
        
        # If not found, try to find in FAMILY CONSTELLATION section
        if not problems_details:
            for i, line in enumerate(self.lines):
                if 'FAMILY CONSTELLATION' in line.upper():
                    # Look for problem-related content in this section
                    for j in range(i+1, min(i+30, len(self.lines))):
                        next_line = self.lines[j].strip()
                        if 'problem' in next_line.lower() and 'area' in next_line.lower():
                            problems_details = next_line.split(':', 1)[1].strip() if ':' in next_line else next_line
                            break
        
        pre_assessment_data = {
            "assessment": "",
            "problems_details": problems_details if problems_details else ""
        }
        
        self.extracted_data['pre_assessment'] = pre_assessment_data
        return pre_assessment_data
    
    def extract_chief_complaints(self) -> Dict[str, Any]:
        """Extract chief complaints data"""
        chief_complaint = ""
        problems_summary = self.extracted_data.get('pre_assessment', {}).get('problems_details', "")
        
        # Look for CHIEF COMPLAINTS section
        for i, line in enumerate(self.lines):
            if 'CHIEF COMPLAINTS:' in line.upper():
                if ':' in line:
                    chief_complaint = line.split(':', 1)[1].strip()
                else:
                    # Check next line
                    if i+1 < len(self.lines):
                        chief_complaint = self.lines[i+1].strip()
                break
        
        # Check for delayed speech and language
        delayed_speech = False
        for line in self.lines:
            if any(term in line.lower() for term in ['delayed speech', 'speech delay', 'language delay', 'not able to speak']):
                delayed_speech = True
                break
        
        chief_complaints_data = {
            "pre_assessment": "",
            "down_syndrome": self.find_field_in_lines(['down syndrome', 'down\'s syndrome']).lower() in ['yes', 'true', 'present'],
            "Chief_complaint": chief_complaint,
            "problems_summary": problems_summary,
            "delayed_speech_language": delayed_speech,
            "poor_attention_sitting": self.find_field_in_lines(['poor attention', 'sitting tolerance']).lower() in ['yes', 'true', 'present'],
            "minimal_vocalization": self.find_field_in_lines(['minimal vocalization', 'vocalization']).lower() in ['yes', 'true', 'present'],
            "toilet_training_needed": self.find_field_in_lines(['toilet training', 'potty training']).lower() in ['yes', 'true', 'present'],
            "behavioral_issues": ""
        }
        
        self.extracted_data['chief_complaints'] = chief_complaints_data
        return chief_complaints_data
    
    def extract_development_milestones(self) -> Dict[str, Any]:
        """Extract development milestones data"""
        milestone_data = {}
        
        # Look for DEVELOPMENTAL MILESTONES section
        found_section = False
        for i, line in enumerate(self.lines):
            if 'DEVELOPMENTAL MILESTONES' in line.upper():
                found_section = True
                # Look for table or list of milestones
                for j in range(i+1, min(i+20, len(self.lines))):
                    current_line = self.lines[j]
                    
                    # Check for sitting milestone
                    if 'sitting' in current_line.lower():
                        if '|' in current_line:  # Table format
                            parts = [p.strip() for p in current_line.split('|')]
                            if len(parts) >= 3:
                                milestone_data['Sitting'] = parts[2].strip() if parts[2].strip() else "___"
                        elif ':' in current_line:
                            parts = current_line.split(':', 1)
                            milestone_data['Sitting'] = parts[1].strip() if len(parts) > 1 else "___"
                    
                    # Check for creeping/crawling milestone
                    elif any(word in current_line.lower() for word in ['creeping', 'crawling']):
                        if '|' in current_line:
                            parts = [p.strip() for p in current_line.split('|')]
                            if len(parts) >= 3:
                                milestone_data['Creeping'] = parts[2].strip() if parts[2].strip() else "___"
                        elif ':' in current_line:
                            parts = current_line.split(':', 1)
                            milestone_data['Creeping'] = parts[1].strip() if len(parts) > 1 else "___"
                    
                    # Check for standing milestone
                    elif 'standing' in current_line.lower():
                        if '|' in current_line:
                            parts = [p.strip() for p in current_line.split('|')]
                            if len(parts) >= 3:
                                milestone_data['Standing'] = parts[2].strip() if parts[2].strip() else "___"
                        elif ':' in current_line:
                            parts = current_line.split(':', 1)
                            milestone_data['Standing'] = parts[1].strip() if len(parts) > 1 else "___"
                    
                    # Check for walking milestone
                    elif 'walking' in current_line.lower():
                        if '|' in current_line:
                            parts = [p.strip() for p in current_line.split('|')]
                            if len(parts) >= 3:
                                milestone_data['Walking'] = parts[2].strip() if parts[2].strip() else "___"
                        elif ':' in current_line:
                            parts = current_line.split(':', 1)
                            milestone_data['Walking'] = parts[1].strip() if len(parts) > 1 else "___"
                    
                    # Check for speaking milestone
                    elif 'speaking' in current_line.lower():
                        if '|' in current_line:
                            parts = [p.strip() for p in current_line.split('|')]
                            if len(parts) >= 3:
                                milestone_data['Speaking'] = parts[2].strip() if parts[2].strip() else "___"
                        elif ':' in current_line:
                            parts = current_line.split(':', 1)
                            milestone_data['Speaking'] = parts[1].strip() if len(parts) > 1 else "___"
                
                break
        
        # Ensure all required milestones are present
        required_milestones = ['Sitting', 'Creeping', 'Standing', 'Walking', 'Speaking']
        for milestone in required_milestones:
            if milestone not in milestone_data:
                milestone_data[milestone] = "___"
        
        development_milestones_data = {
            "milestone_data": milestone_data,
            "pre_assessment": ""
        }
        
        self.extracted_data['development_milestones'] = development_milestones_data
        return development_milestones_data
    
    def extract_comprehensive_exam(self) -> Dict[str, Any]:
        """Extract comprehensive examination data"""
        comprehensive_exam_data = {
            "pre_assessment": "",
            "general_appearance": "",
            "head_shape_size": "",
            "higher_function": "",
            "cognition_perception": "",
            "orientation": "",
            "attention": "",
            "memory": "",
            "spasticity": False,
            "ataxia": False,
            "flaccidity": False,
            "rigidity": False,
            "atonia": False,
            "tremor": False,
            "posture": "",
            "gait": "",
            "upper_limb_power": 0,
            "lower_limb_power": 0,
            "muscle_tone": "",
            "reach_midline": "",
            "reach_across": "",
            "grasp_spherical": "",
            "grasp_cylindrical": "",
            "grasp_hook": "",
            "prehensions_lateral": "",
            "prehensions_palmar": "",
            "prehensions_chuck": "",
            "opponens": "",
            "in_hand_manipulation": "",
            "bilateral_use_of_hand": "",
            "release": "",
            "handedness": "",
            "pain_sensation": "",
            "touch_sensation": "",
            "temperature_sensation": "",
            "proprioception_deep": "",
            "kinesthesia_deep": "",
            "coordination_deep": "",
            "equilibrium_sitting_standing_walking": "",
            "finger_nose": "",
            "finger_finger": "",
            "heal_knee": ""
        }
        
        # Look for ON EXAMINATION section
        found_section = False
        for i, line in enumerate(self.lines):
            if 'ON EXAMINATION' in line.upper() or 'ON EXMINATION' in line.upper():
                found_section = True
                # Process fields in this section
                for j in range(i+1, min(i+50, len(self.lines))):
                    current_line = self.lines[j].strip()
                    
                    if not current_line or (current_line.isupper() and len(current_line) < 100):
                        break
                    
                    # General Appearance
                    if 'General Appearance:' in current_line:
                        value = current_line.split(':', 1)[1].strip() if ':' in current_line else ""
                        if value:
                            comprehensive_exam_data['general_appearance'] = value
                    
                    # Head shape & size
                    elif 'Head shape & size:' in current_line or 'Head shape' in current_line:
                        if ':' in current_line:
                            value = current_line.split(':', 1)[1].strip()
                            comprehensive_exam_data['head_shape_size'] = self.normalize_quality(value)
                    
                    # Higher function
                    elif 'Higher function:' in current_line:
                        if ':' in current_line:
                            value = current_line.split(':', 1)[1].strip()
                            comprehensive_exam_data['higher_function'] = self.normalize_quality(value)
                    
                    # Cognition/perception
                    elif 'Cognition/perception' in current_line or 'Cognition' in current_line:
                        if ':' in current_line:
                            value = current_line.split(':', 1)[1].strip()
                            comprehensive_exam_data['cognition_perception'] = self.normalize_quality(value)
                    
                    # Orientation
                    elif 'Orientation:' in current_line:
                        if ':' in current_line:
                            value = current_line.split(':', 1)[1].strip()
                            comprehensive_exam_data['orientation'] = self.normalize_quality(value)
                    
                    # Attention
                    elif 'Attention:' in current_line:
                        if ':' in current_line:
                            value = current_line.split(':', 1)[1].strip()
                            comprehensive_exam_data['attention'] = self.normalize_quality(value)
                    
                    # Memory
                    elif 'Memory:' in current_line:
                        if ':' in current_line:
                            value = current_line.split(':', 1)[1].strip()
                            comprehensive_exam_data['memory'] = self.normalize_quality(value)
                    
                    # Posture
                    elif 'Posture (' in current_line or 'Posture:' in current_line:
                        if ':' in current_line:
                            value = current_line.split(':', 1)[1].strip()
                        else:
                            value = current_line.replace('Posture', '').strip('() ')
                        if value:
                            comprehensive_exam_data['posture'] = value
                    
                    # Gait
                    elif 'Gait (' in current_line or 'Gait:' in current_line:
                        if ':' in current_line:
                            value = current_line.split(':', 1)[1].strip()
                        else:
                            value = current_line.replace('Gait', '').strip('() ')
                        if value:
                            comprehensive_exam_data['gait'] = value
                    
                    # Muscle power - Upper Limb
                    elif 'Upper Limb:' in current_line or 'Upper limb:' in current_line:
                        if ':' in current_line:
                            value = current_line.split(':', 1)[1].strip()
                            if value and value != 'N':
                                try:
                                    comprehensive_exam_data['upper_limb_power'] = int(re.search(r'\d+', value).group())
                                except:
                                    pass
                    
                    # Muscle power - Lower limb
                    elif 'Lower limb:' in current_line or 'Lower Limb:' in current_line:
                        if ':' in current_line:
                            value = current_line.split(':', 1)[1].strip()
                            if value and value != 'N':
                                try:
                                    comprehensive_exam_data['lower_limb_power'] = int(re.search(r'\d+', value).group())
                                except:
                                    pass
        
        # Set default values for normalized fields if empty
        for field in ['head_shape_size', 'higher_function', 'cognition_perception', 'orientation', 'attention', 'memory']:
            if not comprehensive_exam_data[field]:
                comprehensive_exam_data[field] = self.normalize_quality("")
        
        self.extracted_data['comprehensive_exam'] = comprehensive_exam_data
        return comprehensive_exam_data
    
    def extract_parent_expectations(self) -> Dict[str, Any]:
        """Extract parent expectations data"""
        parent_message = ""
        therapist_recommendation = ""
        
        # Look for Parents Expectations
        for i, line in enumerate(self.lines):
            if 'Parents Expectations:' in line or 'Parent Expectations:' in line or 'PARENTAL EXPECTATIONS' in line.upper():
                if ':' in line:
                    parent_message = line.split(':', 1)[1].strip()
                else:
                    # Check next lines
                    for j in range(i+1, min(i+10, len(self.lines))):
                        next_line = self.lines[j].strip()
                        if not next_line or 'Therapist' in next_line or 'THERAPY' in next_line.upper():
                            break
                        parent_message += " " + next_line
                    parent_message = parent_message.strip()
                break
        
        # Look for Therapist's Opinion/Recommendation
        for i, line in enumerate(self.lines):
            if any(term in line for term in ["Therapist's Opinion", "Therapist's Recommendation", "Therapist Opinion", "THERAPY/OPINION/RECOMMENDATION"]):
                if ':' in line:
                    therapist_recommendation = line.split(':', 1)[1].strip()
                else:
                    # Check next lines
                    for j in range(i+1, min(i+10, len(self.lines))):
                        next_line = self.lines[j].strip()
                        if not next_line or 'Signature' in next_line or 'SIGNATURE' in next_line:
                            break
                        therapist_recommendation += " " + next_line
                    therapist_recommendation = therapist_recommendation.strip()
                break
        
        parent_expectations_data = {
            "pre_assessment": "",
            "parent_message": parent_message,
            "therapist_recommendation": therapist_recommendation
        }
        
        self.extracted_data['parent_expectations'] = parent_expectations_data
        return parent_expectations_data
    
    def extract_speech_language_assessment(self) -> Dict[str, Any]:
        """Extract speech and language assessment data"""
        diagnostic_impression = ""
        provisional_diagnosis = ""
        hearing_loss_degree = ""
        hearing_reported_by = ""
        hearing_report_date = ""
        speech_therapist = ""
        recommendations = ""
        
        # Look for DIAGNOSTIC FORMULATION & CLINICAL IMPRESSION
        for i, line in enumerate(self.lines):
            if 'DIAGNOSTIC FORMULATION & CLINICAL IMPRESSION' in line.upper():
                # Get content after the heading
                content_lines = []
                for j in range(i+1, min(i+15, len(self.lines))):
                    next_line = self.lines[j].strip()
                    if not next_line or 'PROVISIONAL DIAGNOSIS' in next_line.upper():
                        break
                    content_lines.append(next_line)
                diagnostic_impression = ' '.join(content_lines)
                break
        
        # Look for PROVISIONAL DIAGNOSIS
        for i, line in enumerate(self.lines):
            if 'PROVISIONAL DIAGNOSIS' in line.upper():
                # Get content after the heading
                content_lines = []
                for j in range(i+1, min(i+10, len(self.lines))):
                    next_line = self.lines[j].strip()
                    if not next_line or 'RECOMMENDATIONS' in next_line.upper():
                        break
                    content_lines.append(next_line)
                provisional_diagnosis = ' '.join(content_lines)
                break
        
        # Look for Degree of Hearing Loss
        for i, line in enumerate(self.lines):
            if 'Degree of Hearing Loss' in line or 'Degree of hearing loss' in line:
                if ':' in line:
                    value = line.split(':', 1)[1].strip()
                    if value and value.lower() not in ['n/a', 'na']:
                        hearing_loss_degree = value
                
                # Look for reported by and date in next lines
                for j in range(i+1, min(i+5, len(self.lines))):
                    next_line = self.lines[j]
                    if 'reported by' in next_line.lower():
                        if ':' in next_line:
                            reported_by = next_line.split(':', 1)[1].strip()
                            hearing_reported_by = reported_by
                    
                    # Try to extract date
                    date_match = re.search(r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}', next_line)
                    if date_match:
                        hearing_report_date = self.parse_date_from_text(date_match.group(0))
                break
        
        # Look for Speech Therapist signature
        for i, line in enumerate(self.lines):
            if 'SPEECH THERAPIST' in line.upper() or 'Speech Therapist' in line:
                # Extract name after colon or in next line
                if ':' in line:
                    name = line.split(':', 1)[1].strip()
                    speech_therapist = name
                elif i+1 < len(self.lines):
                    speech_therapist = self.lines[i+1].strip()
                break
        
        # Look for Recommendations in speech section
        for i, line in enumerate(self.lines):
            if 'RECOMMENDATIONS' in line.upper():
                # Check context - look back a few lines for speech/language context
                context_lines = self.lines[max(0, i-5):i]
                context_text = ' '.join(context_lines).lower()
                if any(term in context_text for term in ['speech', 'language', 'diagnostic', 'provisional']):
                    content_lines = []
                    for j in range(i+1, min(i+15, len(self.lines))):
                        next_line = self.lines[j].strip()
                        if not next_line or 'SIGNATURE' in next_line.upper():
                            break
                        content_lines.append(next_line)
                    recommendations = ' '.join(content_lines)
                    break
        
        # Rate of speech
        rate_of_speech = ""
        for line in self.lines:
            if 'RATE OF SPEECH:' in line.upper():
                if ':' in line:
                    rate_of_speech = line.split(':', 1)[1].strip()
                break
        
        # Uses hearing aid
        uses_hearing_aid = False
        for line in self.lines:
            if 'Uses hearing aid' in line or 'hearing aid' in line.lower() and 'yes' in line.lower():
                uses_hearing_aid = True
                break
        
        # Default date if not found
        if not hearing_report_date:
            hearing_report_date = "2025-02-12"
        
        speech_assessment_data = {
            "pre_assessment": "",
            "diagnostic_impression": diagnostic_impression,
            "provisional_diagnosis": provisional_diagnosis,
            "hearing_loss_degree": hearing_loss_degree,
            "hearing_reported_by": hearing_reported_by,
            "hearing_report_date": hearing_report_date,
            "uses_hearing_aid": uses_hearing_aid,
            "rate_of_speech": rate_of_speech if rate_of_speech else "[No information available]",
            "speech_therapist": speech_therapist,
            "recommendations": recommendations
        }
        
        self.extracted_data['speech_language_assessment'] = speech_assessment_data
        return speech_assessment_data
    
    def extract_speech_milestones(self) -> List[Dict[str, Any]]:
        """Extract speech milestones data"""
        milestone_types = [
            {"milestone": "Vocalization", "normal_age": "0-3 months"},
            {"milestone": "Babbling", "normal_age": "4-6 months"},
            {"milestone": "First word", "normal_age": "1 year"},
            {"milestone": "Two words", "normal_age": "9-18 months"},
            {"milestone": "Phrases", "normal_age": "2 years"},
            {"milestone": "Sentences", "normal_age": "2 ½ years"}
        ]
        
        speech_milestones_data = []

        # Look for SPEECH DEVELOPMENTAL MILESTONES table
        for i, line in enumerate(self.lines):
            if 'SPEECH DEVELOPMENTAL MILESTONES' in line.upper():
                # Look for table rows
                for j in range(i+1, min(i+20, len(self.lines))):
                    table_line = self.lines[j]
                    if '|' in table_line:
                        parts = [p.strip() for p in table_line.split('|')]
                        if len(parts) >= 4:
                            try:
                                milestone_num = int(parts[0])
                                if 1 <= milestone_num <= 6:
                                    milestone_info = milestone_types[milestone_num-1]
                                    milestone_entry = {
                                        "speech_assessment": "",
                                        "milestone": milestone_info["milestone"],
                                        "normal_age": milestone_info["normal_age"],
                                        "delayed_by": parts[3] if len(parts) > 3 else "",
                                        "status": ""
                                    }
                                    
                                    # Determine status
                                    delayed_by = milestone_entry["delayed_by"].lower()
                                    if "normal" in delayed_by:
                                        milestone_entry["status"] = "present"
                                    elif "not developed" in delayed_by or "not yet" in delayed_by:
                                        milestone_entry["status"] = "absent"
                                    elif "delayed" in delayed_by:
                                        milestone_entry["status"] = "delayed"
                                    elif delayed_by:
                                        milestone_entry["status"] = "present"
                                    else:
                                        milestone_entry["status"] = "absent"
                                    
                                    speech_milestones_data.append(milestone_entry)
                            except:
                                continue
        
        # If table not found, use default entries
        if not speech_milestones_data:
            for milestone_info in milestone_types:
                milestone_entry = {
                    "speech_assessment": "",
                    "milestone": milestone_info["milestone"],
                    "normal_age": milestone_info["normal_age"],
                    "delayed_by": "",
                    "status": "absent"
                }
                speech_milestones_data.append(milestone_entry)

        self.extracted_data['speech_milestones'] = speech_milestones_data
        return speech_milestones_data
    
    def extract_speech_mechanisms(self) -> Dict[str, Any]:
        """Extract speech mechanisms data"""
        speech_mechanisms_data = {
            "speech_assessment": "",
            "lips_appearance": "normal",
            "lips_functional": "normal",
            "teeth_appearance": "normal",
            "teeth_functional": "normal",
            "tongue_appearance": "normal",
            "tongue_functional": "normal",
            "hard_palate_appearance": "normal",
            "hard_palate_functional": "normal",
            "soft_palate_appearance": "normal",
            "soft_palate_functional": "normal",
            "uvula_jaw_appearance": "normal",
            "uvula_jaw_functional": "normal"
        }
        
        self.extracted_data['speech_mechanisms'] = speech_mechanisms_data
        return speech_mechanisms_data
    
    def extract_phonatic_voice_assessment(self) -> Dict[str, Any]:
        """Extract phonatic voice assessment data"""
        consonant = ""
        vowel = ""
        blends = ""
        diphthong_artic = ""
        
        # Look for PHONOLOGICAL/ARTICULATION ASSESSMENT section
        for i, line in enumerate(self.lines):
            if 'PHONOLOGICAL/ARTICULATION ASSESSMENT' in line.upper() or 'ARTICULATION ASSESSMENT' in line.upper():
                # Look for specific fields in next lines
                for j in range(i+1, min(i+15, len(self.lines))):
                    current_line = self.lines[j]
                    
                    if 'Consonant:' in current_line:
                        if ':' in current_line:
                            consonant = current_line.split(':', 1)[1].strip()
                    
                    elif 'Vowel:' in current_line:
                        if ':' in current_line:
                            vowel = current_line.split(':', 1)[1].strip()
                    
                    elif 'Blends:' in current_line:
                        if ':' in current_line:
                            blends = current_line.split(':', 1)[1].strip()
                    
                    elif 'Diphthong:' in current_line:
                        if ':' in current_line:
                            diphthong_artic = current_line.split(':', 1)[1].strip()
        
        # Voice parameters
        loudness = ""
        quality = ""
        pitch = ""
        
        for line in self.lines:
            if 'Loudness:' in line:
                if ':' in line:
                    loudness = line.split(':', 1)[1].strip()
            elif 'Quality:' in line:
                if ':' in line:
                    quality = line.split(':', 1)[1].strip()
            elif 'Pitch:' in line:
                if ':' in line:
                    pitch = line.split(':', 1)[1].strip()
        
        phonatic_voice_data = {
            "speech_assessment": "",
            "consonant": consonant,
            "vowel": vowel,
            "blends": blends,
            "diphthong_artic": diphthong_artic,
            "loudness": loudness,
            "quality": quality,
            "pitch": pitch,
            "diphthong_voice": diphthong_artic
        }
        
        self.extracted_data['phonatic_voice_assessment'] = phonatic_voice_data
        return phonatic_voice_data
    
    def extract_special_education_assessment(self) -> Dict[str, Any]:
        """Extract special education assessment data"""
        conclusion = ""
        for line in self.lines:
            if 'Conclusion:' in line or 'CONCLUSION:' in line:
                if ':' in line:
                    conclusion = line.split(':', 1)[1].strip()
                break
        
        # Get recommendations from speech assessment
        recommendations = self.extracted_data.get('speech_language_assessment', {}).get('recommendations', "")
        
        # Placement recommendation
        placement_recommendation = ""
        for line in self.lines:
            if 'placement' in line.lower() and 'mbcn' in line.lower():
                placement_recommendation = line
                break
        
        # Extract names from signatures
        special_educator_name = ""
        director_principal_name = ""
        vice_principal_name = ""
        parent_guardian_name = ""
        
        # Look for signature section
        for i, line in enumerate(self.lines):
            line_stripped = line.strip()
            if 'Signature:' in line_stripped or 'SIGNATURE:' in line_stripped:
                # Check context to determine whose signature
                context = ' '.join(self.lines[max(0, i-3):i+1]).lower()
                
                if 'special educator' in context:
                    if ':' in line_stripped:
                        name = line_stripped.split(':', 1)[1].strip()
                        special_educator_name = name
                
                elif 'principal' in context or 'superintendent' in context:
                    if ':' in line_stripped:
                        name = line_stripped.split(':', 1)[1].strip()
                        director_principal_name = name
                
                elif 'vice principal' in context:
                    if ':' in line_stripped:
                        name = line_stripped.split(':', 1)[1].strip()
                        vice_principal_name = name
                
                elif 'parent' in context or 'guardian' in context:
                    if ':' in line_stripped:
                        name = line_stripped.split(':', 1)[1].strip()
                        parent_guardian_name = name
        
        special_education_data = {
            "pre_assessment": "",
            "conclusion": conclusion,
            "recommendations": recommendations,
            "placement_recommendation": placement_recommendation,
            "special_educator_name": special_educator_name,
            "special_educator_date": "2025-02-15",
            "director_principal_name": director_principal_name,
            "director_principal_date": "2025-02-16",
            "vice_principal_name": vice_principal_name,
            "vice_principal_date": "2025-02-16",
            "parent_guardian_name": parent_guardian_name,
            "parent_guardian_date": "2025-02-17"
        }
        
        self.extracted_data['special_education_assessment'] = special_education_data
        return special_education_data
    
    def normalize_self_help_value(self, value: str) -> str:
        """Normalize self-help values"""
        if not value:
            return "needs_help"
        
        value_lower = value.lower()
        
        if any(word in value_lower for word in ['independent', 'independently']):
            return "independent"
        elif any(word in value_lower for word in ['dependent', 'depend', 'requires', 'unable']):
            return "dependent"
        else:
            return "needs_help"

    def extract_self_help_skills(self) -> Dict[str, Any]:
        """Extract self-help skills data"""
        # Extract values
        eating = ""
        drinking = ""
        toilet_habits = ""
        brushing = ""
        bathing = ""
        dressing_undressing = ""
        wearing_footwear = ""
        grooming = ""
        
        # Look for SELF-HELP SKILLS section
        for i, line in enumerate(self.lines):
            if 'SELF-HELP SKILLS:' in line.upper():
                # Check next lines for values
                for j in range(i+1, min(i+20, len(self.lines))):
                    current_line = self.lines[j]
                    
                    if 'Eating:' in current_line:
                        if ':' in current_line:
                            eating = current_line.split(':', 1)[1].strip()
                    
                    elif 'Drinking:' in current_line:
                        if ':' in current_line:
                            drinking = current_line.split(':', 1)[1].strip()
                    
                    elif 'Toileting:' in current_line or 'Toilet habits:' in current_line:
                        if ':' in current_line:
                            toilet_habits = current_line.split(':', 1)[1].strip()
                    
                    elif 'Brushing:' in current_line:
                        if ':' in current_line:
                            brushing = current_line.split(':', 1)[1].strip()
                    
                    elif 'Bathing:' in current_line:
                        if ':' in current_line:
                            bathing = current_line.split(':', 1)[1].strip()
                    
                    elif 'Dressing:' in current_line or 'Doffing and Donning:' in current_line:
                        if ':' in current_line:
                            dressing_undressing = current_line.split(':', 1)[1].strip()
                    
                    elif 'Wearing shoes:' in current_line or 'Wearing footwear:' in current_line:
                        if ':' in current_line:
                            wearing_footwear = current_line.split(':', 1)[1].strip()
                    
                    elif 'Grooming:' in current_line:
                        if ':' in current_line:
                            grooming = current_line.split(':', 1)[1].strip()
        
        self_help_skills_data = {
            "special_education": "",
            "eating": self.normalize_self_help_value(eating),
            "drinking": self.normalize_self_help_value(drinking),
            "toilet_habits": self.normalize_self_help_value(toilet_habits),
            "brushing": self.normalize_self_help_value(brushing),
            "bathing": self.normalize_self_help_value(bathing),
            "dressing_undressing": self.normalize_self_help_value(dressing_undressing),
            "wearing_footwear": self.normalize_self_help_value(wearing_footwear),
            "grooming": self.normalize_self_help_value(grooming)
        }
        
        self.extracted_data['self_help_skills'] = self_help_skills_data
        return self_help_skills_data
    
    def extract_socialization(self) -> Dict[str, Any]:
        """Extract socialization data"""
        def get_answer(field_text):
            for line in self.lines:
                if field_text.lower() in line.lower():
                    if ':' in line:
                        return line.split(':', 1)[1].strip()
                    else:
                        return line.split(field_text, 1)[1].strip() if field_text in line else ""
            return ""
        
        socialization_data = {
            "special_education": "",
            "eye_contact": get_answer('she regards face and make eye contact'),
            "expresses_needs": get_answer('expresses his/her needs'),
            "shows_preferences": get_answer('she shows likes and disliked'),
            "shows_emotions": get_answer('she shows a wide variety of emotion'),
            "greets_familiar": get_answer('greets familiar person'),
            "cooperative_play": get_answer('play with 2-3 children cooperatively'),
            "follows_directions": get_answer('can bring object from another room'),
            "neighborhood_mobility": get_answer('neighborhood mobility')
        }
        
        self.extracted_data['socialization'] = socialization_data
        return socialization_data
    
    def extract_cognitive_assessment(self) -> Dict[str, Any]:
        """Extract cognitive assessment data"""
        disability_level = ""
        for line in self.lines:
            if 'level of disability' in line.lower():
                if ':' in line:
                    disability_level = line.split(':', 1)[1].strip()
                break
        
        identifies_objects = ""
        for line in self.lines:
            if 'identification of familiar objects' in line.lower():
                if ':' in line:
                    identifies_objects = line.split(':', 1)[1].strip()
                break
        
        uses_objects = ""
        for line in self.lines:
            if 'use of familiar objects' in line.lower():
                if ':' in line:
                    uses_objects = line.split(':', 1)[1].strip()
                break
        
        def get_bool_field(field_text):
            for line in self.lines:
                if field_text.lower() in line.lower():
                    line_lower = line.lower()
                    if 'yes' in line_lower or 'true' in line_lower or 'present' in line_lower:
                        return True
            return False
        
        cognitive_assessment_data = {
            "special_education": "",
            "disability_level": disability_level,
            "disability_reported_by": self.find_field_in_lines(['reported by']),
            "disability_assessment_date": "2025-12-15",
            "identifies_objects": identifies_objects,
            "uses_objects": uses_objects,
            "follows_instructions": self.find_field_in_lines(['following simple instructions', 'follows instructions']),
            "danger_awareness": self.find_field_in_lines(['awareness of dangers']),
            "color_concept": get_bool_field('color concept'),
            "size_concept": get_bool_field('size concept'),
            "sex_concept": get_bool_field('sex concept'),
            "shape_concept": get_bool_field('shape concept'),
            "number_concept": get_bool_field('number concept'),
            "time_concept": get_bool_field('time concept'),
            "money_concept": get_bool_field('money concept')
        }
        
        self.extracted_data['cognitive_assessment'] = cognitive_assessment_data
        return cognitive_assessment_data
    
    def extract_other_assessments(self) -> Dict[str, Any]:
        """Extract other assessments data"""
        def get_field_value(field_text, default=""):
            for line in self.lines:
                if field_text.lower() in line.lower():
                    if ':' in line:
                        return line.split(':', 1)[1].strip()
            return default
        
        # Extract likes and dislikes
        likes_dislikes = ""
        for i, line in enumerate(self.lines):
            if 'likes and dislikes' in line.lower() or 'likes / dislikes' in line.lower():
                if ':' in line:
                    likes_dislikes = line.split(':', 1)[1].strip()
                else:
                    # Check next line
                    if i+1 < len(self.lines):
                        likes_dislikes = self.lines[i+1].strip()
                break
        
        # Extract special abilities
        special_abilities = ""
        for i, line in enumerate(self.lines):
            if 'special ability' in line.lower() or 'special abilities' in line.lower():
                if ':' in line:
                    special_abilities = line.split(':', 1)[1].strip()
                else:
                    # Check next line
                    if i+1 < len(self.lines):
                        special_abilities = self.lines[i+1].strip()
                break
        
        other_assessments_data = {
            "special_education": "",
            "academic_exposure": get_field_value('previous education', ''),
            "reading_ability": get_field_value('reading', ''),
            "writing_ability": get_field_value('writing', ''),
            "arithmetic_ability": get_field_value('arithmetic', ''),
            "sensory_status": get_field_value('sensory status', ''),
            "behavioral_observation": "",
            "attention_span": get_field_value('attention span', ''),
            "cooperation_level": get_field_value('child\'s co-operation', ''),
            "reported_problem_behavior": get_field_value('problem behavior reported by parents', ''),
            "observed_problem_behavior": get_field_value('problem behavior observed by Educator', ''),
            "child_father_relationship": get_field_value('child father relationship', ''),
            "child_mother_relationship": get_field_value('child mother relationship', ''),
            "child_siblings_relationship": get_field_value('child siblings relationship', ''),
            "child_other_members_relationship": get_field_value('child other members', ''),
            "father_mother_relationship": get_field_value('father mother', ''),
            "among_other_members_relationship": get_field_value('among other members', ''),
            "likes_dislikes": likes_dislikes,
            "special_abilities": special_abilities
        }
        
        self.extracted_data['other_assessments'] = other_assessments_data
        return other_assessments_data
    
    def extract_basic_observation(self) -> Dict[str, Any]:
        """Extract basic observation pre-assessment data"""
        def get_yes_no_field(field_text, default="no"):
            for line in self.lines:
                if field_text.lower() in line.lower():
                    line_lower = line.lower()
                    if 'yes' in line_lower:
                        return "yes"
                    elif 'no' in line_lower:
                        return "no"
            return default
        
        basic_observation_data = {
            "assessment": "",
            "assessment_date": "2025-12-17",
            "socialization": self.find_field_in_lines(['socialization level']) or "poor",
            "eye_contact": self.find_field_in_lines(['eye contact behavior']) or "poor",
            "peers_group_relationship": self.find_field_in_lines(['peers group relationship']) or "poor",
            "instructions_following": self.find_field_in_lines(['instructions following']) or "fair",
            "auditory_sound": {
                "covers_ears": get_yes_no_field('covers ears'),
                "difficulty_with_background_sounds": get_yes_no_field('difficulty with background sounds'),
                "does_not_respond_to_name": get_yes_no_field('does not respond to name')
            },
            "visual_sight": {
                "happy_in_dark": get_yes_no_field('happy in dark'),
                "covers_eyes_from_light": get_yes_no_field('covers eyes from light'),
                "stares_at_moving_objects": get_yes_no_field('stares at moving objects')
            },
            "gustatory_taste": {
                "picky_eater": get_yes_no_field('picky eater'),
                "chews_or_licks_non_food_objects": get_yes_no_field('chews or licks non food objects'),
                "mouthing_objects_or_fingers": get_yes_no_field('mouthing objects or fingers', "yes")
            },
            "olfactory_smell": {
                "smells_non_food_objects_or_people": get_yes_no_field('smells non food objects or people'),
                "strong_preference_for_certain_smells": get_yes_no_field('strong preference for certain smells')
            },
            "tactile_touch": {
                "avoids_barefoot_on_uneven_surfaces": get_yes_no_field('avoids barefoot on uneven surfaces'),
                "reacts_emotionally_or_aggressively_to_touch": get_yes_no_field('reacts emotionally or aggressively to touch'),
                "withdraws_from_splashing_water": get_yes_no_field('withdraws from splashing water'),
                "difficulty_standing_close_to_others": get_yes_no_field('difficulty standing close to others'),
                "rubs_or_scratches_touched_area": get_yes_no_field('rubs or scratches touched area', "not_observed")
            }
        }
        
        self.extracted_data['basic_observation'] = basic_observation_data
        return basic_observation_data
    
    def extract_sensory_assessment(self) -> Dict[str, Any]:
        """Extract sensory pre-assessment data"""
        def get_field_value(field_text, default=""):
            for line in self.lines:
                if field_text.lower() in line.lower():
                    if ':' in line:
                        return line.split(':', 1)[1].strip()
            return default
        
        def get_yes_no_field(field_text, default="no"):
            for line in self.lines:
                if field_text.lower() in line.lower():
                    line_lower = line.lower()
                    if 'yes' in line_lower:
                        return "yes"
                    elif 'no' in line_lower:
                        return "no"
            return default
        
        sensory_assessment_data = {
            "assessment": "",
            "assessment_date": "2025-12-17",
            "vestibular_body_movement": {
                "avoids_or_seeks_moving_activities": get_field_value('avoids or seeks moving activities', "seeks moving activities"),
                "takes_unconscious_risks": get_yes_no_field('takes unconscious risks'),
                "twirls_or_spins_self_frequently": get_yes_no_field('twirls or spins self frequently'),
                "dislikes_riding_in_car": get_field_value('dislikes riding in car', "")
            },
            "proprioceptive_body_awareness": {
                "enjoys_falling": get_yes_no_field('enjoys falling'),
                "takes_excessive_risks_during_play": get_yes_no_field('takes excessive risks during play'),
                "turns_whole_body_to_look": get_field_value('turns whole body to look', ""),
                "prefers_w_sitting_or_supported_sitting": get_yes_no_field('prefers w sitting or supported sitting'),
                "jumping": get_yes_no_field('jumping', "yes")
            },
            "interoception": {
                "basic_concepts_of_hunger": get_field_value('basic concepts of hunger', "intact"),
                "basic_concepts_of_toilet": get_field_value('basic concepts of toilet', "poor")
            },
            "behavior_issues": {
                "hitting_self_or_others": get_yes_no_field('hitting self or others'),
                "pinching_self_or_others": get_yes_no_field('pinching self or others'),
                "biting_self_or_others": get_yes_no_field('biting self or others'),
                "making_different_types_of_sounds": get_yes_no_field('making different types of sounds', "yes")
            },
            "object_throwing": {
                "irrelevant_laughing": get_yes_no_field('irrelevant laughing'),
                "self_talking_or_echolalia": get_yes_no_field('self talking or echolalia'),
                "throws_objects": get_yes_no_field('throws objects', "yes")
            },
            "self_stimulating_behavior": {
                "hand_flapping": get_yes_no_field('hand flapping'),
                "finger_flickering": get_yes_no_field('finger flickering'),
                "jumping_or_rocking_when_happy_or_anxious": get_yes_no_field('jumping or rocking when happy or anxious', "yes"),
                "attachment_with_certain_objects": get_yes_no_field('attachment with certain objects')
            }
        }
        
        self.extracted_data['sensory_assessment'] = sensory_assessment_data
        return sensory_assessment_data
    
    def extract_all_data(self):
        """Extract all data from the document"""
        self.extract_assessment_data()
        self.extract_pre_assessment_data()
        self.extract_chief_complaints()
        self.extract_development_milestones()
        self.extract_comprehensive_exam()
        self.extract_parent_expectations()
        self.extract_speech_language_assessment()
        self.extract_speech_milestones()
        self.extract_speech_mechanisms()
        self.extract_phonatic_voice_assessment()
        self.extract_special_education_assessment()
        self.extract_self_help_skills()
        self.extract_socialization()
        self.extract_cognitive_assessment()
        self.extract_other_assessments()
        self.extract_basic_observation()
        self.extract_sensory_assessment()
        return self.extracted_data


def make_api_request(endpoint, data, bearer_token=None):
    """Make API request to the specified endpoint"""
    try:
        headers = {'Content-Type': 'application/json'}
        if bearer_token and bearer_token.strip():
            headers['Authorization'] = f'Bearer {bearer_token}'
            
        response = requests.post(endpoint, json=data, headers=headers)
        if response.status_code in [200, 201]:
            return True, response.json()
        else:
            return False, f"Error {response.status_code}: {response.text}"
    except Exception as e:
        return False, str(e)


def extract_id_from_response(result, entity_type=""):
    """Extract ID from API response"""
    if not result:
        return None
    
    possible_id_fields = ['id', f'{entity_type}_id', 'pk', 'uid', 'ID', 'Id', 'uuid']
    
    if isinstance(result, (int, str)):
        return str(result)
    
    if isinstance(result, dict):
        for field in possible_id_fields:
            if field in result and result[field] is not None:
                return str(result[field])
        
        if 'data' in result and isinstance(result['data'], dict):
            for field in possible_id_fields:
                if field in result['data'] and result['data'][field] is not None:
                    return str(result['data'][field])
        
        for key, value in result.items():
            if isinstance(value, (int, str)):
                return str(value)
    
    return None


def main():
    st.set_page_config(page_title="Assessment Data Extractor & API Pusher (IMPROVED)", layout="wide")
    
    st.title("🎯 Assessment Data Extractor & API Pusher (IMPROVED)")
    st.markdown("""
    ✅ **IMPROVED DATA EXTRACTION:**
    1. Better student profile extraction from ASSESSMENT FORM
    2. Fixed D.O.A extraction
    3. Enhanced field extraction based on document structure
    4. Improved normalization of quality fields
    5. Better section boundary detection
    """)
    
    st.subheader("🔐 API Configuration")
    bearer_token = st.text_input(
        "Bearer Token (Optional)", 
        type="password", 
        placeholder="Enter your bearer token here"
    )
    
    st.divider()
    
    uploaded_file = st.file_uploader("Choose a DOCX file", type="docx")
    
    if uploaded_file is not None:
        try:
            # Read file content once
            file_content = uploaded_file.read()
            
            with st.spinner("Extracting data from document..."):
                # Extract student profile data
                student_extractor = StudentProfileExtractor(file_content)
                student_profile = student_extractor.extract_student_profile()
                
                # Extract assessment data
                extractor = AssessmentDataExtractor(file_content)
                extracted_data = extractor.extract_all_data()
                
                # Add student profile to extracted data
                extracted_data['student_profile'] = student_profile
            
            st.success("✅ Data extraction completed successfully!")
            
            # Initialize session state for IDs
            if 'student_id' not in st.session_state:
                st.session_state.student_id = None
            if 'assessment_id' not in st.session_state:
                st.session_state.assessment_id = None
            if 'pre_assessment_id' not in st.session_state:
                st.session_state.pre_assessment_id = None
            if 'speech_assessment_id' not in st.session_state:
                st.session_state.speech_assessment_id = None
            if 'special_education_id' not in st.session_state:
                st.session_state.special_education_id = None
            
            with st.expander("📄 View Document Text", expanded=False):
                st.text_area("Full Document Text:", value=extractor.full_text, height=300)
            
            with st.expander("🔍 View All Extracted Fields (JSON)", expanded=True):
                st.json(extracted_data)
            
            st.divider()
            
            # ============================================================
            # OPTIONAL STUDENT PROFILE CREATION BLOCK
            # ============================================================
            st.subheader("👤 Optional: Create Student Profile")
            st.info("""
            **Use this only if student ID is not already generated.** 
            If student already exists in the system, you can skip this section 
            and enter the existing student ID in the Assessment section below.
            """)
            
            col1, col2 = st.columns([2, 1])
            
            with col1:
                with st.expander("View Student Profile Data", expanded=False):
                    st.json(extracted_data.get('student_profile', {}))
                
                json_student_data = extracted_data.get('student_profile', {})
                edited_json_student = st.text_area(
                    "Edit Student Profile:",
                    value=json.dumps(json_student_data, indent=2, ensure_ascii=False),
                    height=300,
                    key="json_student_editor"
                )
                
            with col2:
                st.markdown("### Create Student")
                if st.button("Create Student Profile", type="primary", key="create_student"):
                    try:
                        json_student_parsed = json.loads(edited_json_student)
                        success, result = make_api_request(
                            "http://mbcnschool.ai:5000/api/students/",
                            json_student_parsed,
                            bearer_token
                        )
                        
                        if success:
                            st.success("✅ Student profile created successfully!")
                            st.json(result)
                            student_id = extract_id_from_response(result, "student")
                            
                            if student_id:
                                st.session_state.student_id = str(student_id)
                                st.info(f"**Generated Student ID:** `{student_id}`")
                                st.rerun()
                        else:
                            st.error(f"Failed: {result}")
                    except Exception as e:
                        st.error(f"Error: {str(e)}")
                
                if st.session_state.student_id:
                    st.success(f"**✅ Student ID:**\n{st.session_state.student_id}")
                else:
                    st.warning("No student ID generated yet")
            
            st.divider()
            # ============================================================
            # END OF STUDENT PROFILE BLOCK
            # ============================================================
            
            API_ENDPOINTS = {
                "assessment": "http://mbcnschool.ai:5000/api/assessments/",
                "pre_assessment": "http://mbcnschool.ai:5000/api/pre-assessments/",
                "chief_complaints": "http://mbcnschool.ai:5000/api/chief-complaints/",
                "development_milestones": "http://mbcnschool.ai:5000/api/development-milestones/",
                "comprehensive_exam": "http://mbcnschool.ai:5000/api/comprehensive-exams/",
                "parent_expectations": "http://mbcnschool.ai:5000/api/parent-expectations/",
                "speech_language_assessment": "http://mbcnschool.ai:5000/api/speech-language-assessments/",
                "speech_milestones": "http://mbcnschool.ai:5000/api/speech-milestones/",
                "speech_mechanisms": "http://mbcnschool.ai:5000/api/speech-mechanisms/",
                "phonatic_voice_assessment": "http://mbcnschool.ai:5000/api/phonatic-voice-assessment/",
                "special_education_assessment": "http://mbcnschool.ai:5000/api/special-education-assessments/",
                "self_help_skills": "http://mbcnschool.ai:5000/api/self-help-skills/",
                "socialization": "http://mbcnschool.ai:5000/api/socializations/",
                "cognitive_assessment": "http://mbcnschool.ai:5000/api/cognitive-assessments/",
                "other_assessments": "http://mbcnschool.ai:5000/api/other-assessments/",
                "basic_observation": "http://mbcnschool.ai:5000/api/basic-observation-pre-assessments/",
                "sensory_assessment": "http://mbcnschool.ai:5000/api/sensory-pre-assessments/"
            }
            
            st.subheader("📝 Group 1: Assessment & Pre-Assessment")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("**1️⃣ Assessment Data**")
                with st.expander("View Assessment Fields", expanded=False):
                    st.json(extracted_data.get('assessment', {}))
                
                json1_data = extracted_data.get('assessment', {})
                
                # Add student ID if available
                if st.session_state.student_id:
                    json1_data["student"] = st.session_state.student_id
                else:
                    # Allow manual entry of student ID
                    manual_student_id = st.text_input(
                        "Enter Student ID (if not created above):",
                        key="manual_student_id"
                    )
                    if manual_student_id:
                        json1_data["student"] = manual_student_id
                
                edited_json1 = st.text_area(
                    "Edit Assessment:",
                    value=json.dumps(json1_data, indent=2, ensure_ascii=False),
                    height=400,
                    key="json1_editor"
                )
                
                if st.button("Push Assessment", type="primary", key="push_assessment"):
                    try:
                        json1_parsed = json.loads(edited_json1)
                        success, result = make_api_request(
                            API_ENDPOINTS["assessment"],
                            json1_parsed,
                            bearer_token
                        )
                        
                        if success:
                            st.success("✅ Assessment pushed successfully!")
                            st.json(result)
                            assessment_id = extract_id_from_response(result, "assessment")
                            
                            if assessment_id:
                                st.session_state.assessment_id = str(assessment_id)
                                st.rerun()
                        else:
                            st.error(f"Failed: {result}")
                    except Exception as e:
                        st.error(f"Error: {str(e)}")
                
                if st.session_state.assessment_id:
                    st.info(f"**✅ Generated Assessment ID:** `{st.session_state.assessment_id}`")
            
            with col2:
                st.markdown("**2️⃣ Pre-Assessment Data**")
                with st.expander("View Pre-Assessment Fields", expanded=False):
                    st.json(extracted_data.get('pre_assessment', {}))
                
                json2_data = extracted_data.get('pre_assessment', {})
                if st.session_state.assessment_id:
                    json2_data['assessment'] = st.session_state.assessment_id
                
                edited_json2 = st.text_area(
                    "Edit Pre-Assessment:",
                    value=json.dumps(json2_data, indent=2, ensure_ascii=False),
                    height=400,
                    key="json2_editor"
                )
                
                if st.button("Push Pre-Assessment", type="primary", key="push_pre_assessment"):
                    if not st.session_state.assessment_id:
                        st.error("⚠️ Please push Assessment first!")
                    else:
                        try:
                            json2_parsed = json.loads(edited_json2)
                            json2_parsed['assessment'] = st.session_state.assessment_id
                            
                            success, result = make_api_request(
                                API_ENDPOINTS["pre_assessment"],
                                json2_parsed,
                                bearer_token
                            )
                            
                            if success:
                                st.success("✅ Pre-Assessment pushed!")
                                st.json(result)
                                pre_assessment_id = extract_id_from_response(result)
                                
                                if pre_assessment_id:
                                    st.session_state.pre_assessment_id = str(pre_assessment_id)
                                    st.rerun()
                            else:
                                st.error(f"Failed: {result}")
                        except Exception as e:
                            st.error(f"Error: {str(e)}")
                
                if st.session_state.pre_assessment_id:
                    st.info(f"**✅ Generated Pre-Assessment ID:** `{st.session_state.pre_assessment_id}`")
            
            st.divider()
            
            st.subheader("📋 Pre-Assessment Dependent Fields (3-6)")
            st.info("⚠️ These fields require Pre-Assessment ID to be pushed first")
            
            # 3. Chief Complaints
            with st.expander("3️⃣ Chief Complaints", expanded=False):
                json3_data = extracted_data.get('chief_complaints', {}).copy()
                if st.session_state.pre_assessment_id:
                    json3_data['pre_assessment'] = st.session_state.pre_assessment_id
                st.json(json3_data)
                
                edited_json3 = st.text_area(
                    "Edit Chief Complaints:",
                    value=json.dumps(json3_data, indent=2, ensure_ascii=False),
                    height=300,
                    key="json3_editor"
                )
                
                if st.button("Push Chief Complaints", type="primary", key="push_chief_complaints"):
                    if not st.session_state.pre_assessment_id:
                        st.error("⚠️ Please push Pre-Assessment first!")
                    else:
                        try:
                            json3_parsed = json.loads(edited_json3)
                            json3_parsed['pre_assessment'] = st.session_state.pre_assessment_id
                            
                            success, result = make_api_request(
                                API_ENDPOINTS["chief_complaints"],
                                json3_parsed,
                                bearer_token
                            )
                            
                            if success:
                                st.success("✅ Chief Complaints pushed!")
                                st.json(result)
                            else:
                                st.error(f"Failed: {result}")
                        except Exception as e:
                            st.error(f"Error: {str(e)}")
            
            # 4. Development Milestones
            with st.expander("4️⃣ Development Milestones", expanded=False):
                json4_data = extracted_data.get('development_milestones', {}).copy()
                if st.session_state.pre_assessment_id:
                    json4_data['pre_assessment'] = st.session_state.pre_assessment_id
                st.json(json4_data)
                
                edited_json4 = st.text_area(
                    "Edit Development Milestones:",
                    value=json.dumps(json4_data, indent=2, ensure_ascii=False),
                    height=300,
                    key="json4_editor"
                )
                
                if st.button("Push Development Milestones", type="primary", key="push_development_milestones"):
                    if not st.session_state.pre_assessment_id:
                        st.error("⚠️ Please push Pre-Assessment first!")
                    else:
                        try:
                            json4_parsed = json.loads(edited_json4)
                            json4_parsed['pre_assessment'] = st.session_state.pre_assessment_id
                            
                            if 'milestone_data' in json4_parsed and isinstance(json4_parsed['milestone_data'], dict):
                                json4_parsed['milestone_data'] = json.dumps(json4_parsed['milestone_data'], ensure_ascii=False)
                            
                            success, result = make_api_request(
                                API_ENDPOINTS["development_milestones"],
                                json4_parsed,
                                bearer_token
                            )

                            if success:
                                st.success("✅ Development Milestones pushed!")
                                st.json(result)
                            else:
                                st.error(f"Failed: {result}")
                        except Exception as e:
                            st.error(f"Error: {str(e)}")
            
            # 5. Comprehensive Exam
            with st.expander("5️⃣ Comprehensive Exam", expanded=False):
                json5_data = extracted_data.get('comprehensive_exam', {}).copy()
                if st.session_state.pre_assessment_id:
                    json5_data['pre_assessment'] = st.session_state.pre_assessment_id
                st.json(json5_data)
                
                edited_json5 = st.text_area(
                    "Edit Comprehensive Exam:",
                    value=json.dumps(json5_data, indent=2, ensure_ascii=False),
                    height=300,
                    key="json5_editor"
                )
                
                if st.button("Push Comprehensive Exam", type="primary", key="push_comprehensive_exam"):
                    if not st.session_state.pre_assessment_id:
                        st.error("⚠️ Please push Pre-Assessment first!")
                    else:
                        try:
                            json5_parsed = json.loads(edited_json5)
                            json5_parsed['pre_assessment'] = st.session_state.pre_assessment_id
                            
                            success, result = make_api_request(
                                API_ENDPOINTS["comprehensive_exam"],
                                json5_parsed,
                                bearer_token
                            )
                            
                            if success:
                                st.success("✅ Comprehensive Exam pushed!")
                                st.json(result)
                            else:
                                st.error(f"Failed: {result}")
                        except Exception as e:
                            st.error(f"Error: {str(e)}")
            
            # 6. Parent Expectations
            with st.expander("6️⃣ Parent Expectations", expanded=False):
                json6_data = extracted_data.get('parent_expectations', {}).copy()
                if st.session_state.pre_assessment_id:
                    json6_data['pre_assessment'] = st.session_state.pre_assessment_id
                st.json(json6_data)
                
                edited_json6 = st.text_area(
                    "Edit Parent Expectations:",
                    value=json.dumps(json6_data, indent=2, ensure_ascii=False),
                    height=300,
                    key="json6_editor"
                )
                
                if st.button("Push Parent Expectations", type="primary", key="push_parent_expectations"):
                    if not st.session_state.pre_assessment_id:
                        st.error("⚠️ Please push Pre-Assessment first!")
                    else:
                        try:
                            json6_parsed = json.loads(edited_json6)
                            json6_parsed['pre_assessment'] = st.session_state.pre_assessment_id
                            
                            success, result = make_api_request(
                                API_ENDPOINTS["parent_expectations"],
                                json6_parsed,
                                bearer_token
                            )
                            
                            if success:
                                st.success("✅ Parent Expectations pushed!")
                                st.json(result)
                            else:
                                st.error(f"Failed: {result}")
                        except Exception as e:
                            st.error(f"Error: {str(e)}")
            
            st.divider()
            
            st.subheader("🗣️ Speech Assessment Section (7-10)")
            
            # 7. Speech Language Assessment
            with st.expander("7️⃣ Speech Language Assessment", expanded=False):
                json7_data = extracted_data.get('speech_language_assessment', {}).copy()
                if st.session_state.pre_assessment_id:
                    json7_data['pre_assessment'] = st.session_state.pre_assessment_id
                st.json(json7_data)
                
                edited_json7 = st.text_area(
                    "Edit Speech Language Assessment:",
                    value=json.dumps(json7_data, indent=2, ensure_ascii=False),
                    height=300,
                    key="json7_editor"
                )
                
                if st.button("Push Speech Language Assessment", type="primary", key="push_speech_assessment"):
                    if not st.session_state.pre_assessment_id:
                        st.error("⚠️ Please push Pre-Assessment first!")
                    else:
                        try:
                            json7_parsed = json.loads(edited_json7)
                            json7_parsed['pre_assessment'] = st.session_state.pre_assessment_id
                            
                            success, result = make_api_request(
                                API_ENDPOINTS["speech_language_assessment"],
                                json7_parsed,
                                bearer_token
                            )
                            
                            if success:
                                st.success("✅ Speech Language Assessment pushed!")
                                st.json(result)
                                speech_id = extract_id_from_response(result)
                                
                                if speech_id:
                                    st.session_state.speech_assessment_id = str(speech_id)
                                    st.rerun()
                            else:
                                st.error(f"Failed: {result}")
                        except Exception as e:
                            st.error(f"Error: {str(e)}")
                
                if st.session_state.speech_assessment_id:
                    st.info(f"**✅ Generated Speech Assessment ID:** `{st.session_state.speech_assessment_id}`")
            
            st.info("⚠️ Fields 8-10 require Speech Assessment ID")
            
            # 8. Speech Milestones
            with st.expander("8️⃣ Speech Milestones", expanded=False):
                json8_data = extracted_data.get('speech_milestones', [])
                
                json8_data_updated = []
                for milestone in json8_data:
                    milestone_copy = milestone.copy()
                    if st.session_state.speech_assessment_id:
                        milestone_copy['speech_assessment'] = st.session_state.speech_assessment_id
                    json8_data_updated.append(milestone_copy)
                
                st.json(json8_data_updated)
                
                edited_json8 = st.text_area(
                    "Edit Speech Milestones:",
                    value=json.dumps(json8_data_updated, indent=2, ensure_ascii=False),
                    height=300,
                    key="json8_editor"
                )
                
                if st.button("Push Speech Milestones", type="primary", key="push_speech_milestones"):
                    if not st.session_state.speech_assessment_id:
                        st.error("⚠️ Please push Speech Assessment first!")
                    else:
                        try:
                            json8_parsed = json.loads(edited_json8)
                            
                            for milestone in json8_parsed:
                                milestone['speech_assessment'] = st.session_state.speech_assessment_id
                                success, result = make_api_request(
                                    API_ENDPOINTS["speech_milestones"],
                                    milestone,
                                    bearer_token
                                )
                                
                                if not success:
                                    st.error(f"Failed to push milestone: {result}")
                                    break
                            else:
                                st.success("✅ All Speech Milestones pushed!")
                        except Exception as e:
                            st.error(f"Error: {str(e)}")
            
            # 9. Speech Mechanisms
            with st.expander("9️⃣ Speech Mechanisms", expanded=False):
                json9_data = extracted_data.get('speech_mechanisms', {}).copy()
                if st.session_state.speech_assessment_id:
                    json9_data['speech_assessment'] = st.session_state.speech_assessment_id
                st.json(json9_data)
                
                edited_json9 = st.text_area(
                    "Edit Speech Mechanisms:",
                    value=json.dumps(json9_data, indent=2, ensure_ascii=False),
                    height=300,
                    key="json9_editor"
                )
                
                if st.button("Push Speech Mechanisms", type="primary", key="push_speech_mechanisms"):
                    if not st.session_state.speech_assessment_id:
                        st.error("⚠️ Please push Speech Assessment first!")
                    else:
                        try:
                            json9_parsed = json.loads(edited_json9)
                            json9_parsed['speech_assessment'] = st.session_state.speech_assessment_id
                            
                            success, result = make_api_request(
                                API_ENDPOINTS["speech_mechanisms"],
                                json9_parsed,
                                bearer_token
                            )
                            
                            if success:
                                st.success("✅ Speech Mechanisms pushed!")
                                st.json(result)
                            else:
                                st.error(f"Failed: {result}")
                        except Exception as e:
                            st.error(f"Error: {str(e)}")
            
            # 10. Phonatic Voice Assessment
            with st.expander("🔟 Phonatic Voice Assessment", expanded=False):
                json10_data = extracted_data.get('phonatic_voice_assessment', {}).copy()
                if st.session_state.speech_assessment_id:
                    json10_data['speech_assessment'] = st.session_state.speech_assessment_id
                st.json(json10_data)
                
                edited_json10 = st.text_area(
                    "Edit Phonatic Voice Assessment:",
                    value=json.dumps(json10_data, indent=2, ensure_ascii=False),
                    height=300,
                    key="json10_editor"
                )
                
                if st.button("Push Phonatic Voice Assessment", type="primary", key="push_phonatic_voice"):
                    if not st.session_state.speech_assessment_id:
                        st.error("⚠️ Please push Speech Assessment first!")
                    else:
                        try:
                            json10_parsed = json.loads(edited_json10)
                            json10_parsed['speech_assessment'] = st.session_state.speech_assessment_id
                            
                            success, result = make_api_request(
                                API_ENDPOINTS["phonatic_voice_assessment"],
                                json10_parsed,
                                bearer_token
                            )
                            
                            if success:
                                st.success("✅ Phonatic Voice Assessment pushed!")
                                st.json(result)
                            else:
                                st.error(f"Failed: {result}")
                        except Exception as e:
                            st.error(f"Error: {str(e)}")
            
            st.divider()
            
            st.subheader("🎓 Special Education Section (11-15)")
            
            # 11. Special Education Assessment
            with st.expander("1️⃣1️⃣ Special Education Assessment", expanded=False):
                json11_data = extracted_data.get('special_education_assessment', {}).copy()
                if st.session_state.pre_assessment_id:
                    json11_data['pre_assessment'] = st.session_state.pre_assessment_id
                st.json(json11_data)
                
                edited_json11 = st.text_area(
                    "Edit Special Education Assessment:",
                    value=json.dumps(json11_data, indent=2, ensure_ascii=False),
                    height=300,
                    key="json11_editor"
                )
                
                if st.button("Push Special Education Assessment", type="primary", key="push_special_education"):
                    if not st.session_state.pre_assessment_id:
                        st.error("⚠️ Please push Pre-Assessment first!")
                    else:
                        try:
                            json11_parsed = json.loads(edited_json11)
                            json11_parsed['pre_assessment'] = st.session_state.pre_assessment_id
                            
                            success, result = make_api_request(
                                API_ENDPOINTS["special_education_assessment"],
                                json11_parsed,
                                bearer_token
                            )
                            
                            if success:
                                st.success("✅ Special Education Assessment pushed!")
                                st.json(result)
                                special_ed_id = extract_id_from_response(result)
                                
                                if special_ed_id:
                                    st.session_state.special_education_id = str(special_ed_id)
                                    st.rerun()
                            else:
                                st.error(f"Failed: {result}")
                        except Exception as e:
                            st.error(f"Error: {str(e)}")
                
                if st.session_state.special_education_id:
                    st.info(f"**✅ Generated Special Education ID:** `{st.session_state.special_education_id}`")
            
            st.info("⚠️ Fields 12-15 require Special Education ID")
            
            # 12. Self Help Skills
            with st.expander("1️⃣2️⃣ Self Help Skills", expanded=False):
                json12_data = extracted_data.get('self_help_skills', {}).copy()
                if st.session_state.special_education_id:
                    json12_data['special_education'] = st.session_state.special_education_id
                st.json(json12_data)
                
                edited_json12 = st.text_area(
                    "Edit Self Help Skills:",
                    value=json.dumps(json12_data, indent=2, ensure_ascii=False),
                    height=300,
                    key="json12_editor"
                )
                
                if st.button("Push Self Help Skills", type="primary", key="push_self_help"):
                    if not st.session_state.special_education_id:
                        st.error("⚠️ Please push Special Education Assessment first!")
                    else:
                        try:
                            json12_parsed = json.loads(edited_json12)
                            json12_parsed['special_education'] = st.session_state.special_education_id
                            
                            success, result = make_api_request(
                                API_ENDPOINTS["self_help_skills"],
                                json12_parsed,
                                bearer_token
                            )
                            
                            if success:
                                st.success("✅ Self Help Skills pushed!")
                                st.json(result)
                            else:
                                st.error(f"Failed: {result}")
                        except Exception as e:
                            st.error(f"Error: {str(e)}")
            
            # 13. Socialization
            with st.expander("1️⃣3️⃣ Socialization", expanded=False):
                json13_data = extracted_data.get('socialization', {}).copy()
                if st.session_state.special_education_id:
                    json13_data['special_education'] = st.session_state.special_education_id
                st.json(json13_data)
                
                edited_json13 = st.text_area(
                    "Edit Socialization:",
                    value=json.dumps(json13_data, indent=2, ensure_ascii=False),
                    height=300,
                    key="json13_editor"
                )
                
                if st.button("Push Socialization", type="primary", key="push_socialization"):
                    if not st.session_state.special_education_id:
                        st.error("⚠️ Please push Special Education Assessment first!")
                    else:
                        try:
                            json13_parsed = json.loads(edited_json13)
                            json13_parsed['special_education'] = st.session_state.special_education_id
                            
                            success, result = make_api_request(
                                API_ENDPOINTS["socialization"],
                                json13_parsed,
                                bearer_token
                            )
                            
                            if success:
                                st.success("✅ Socialization pushed!")
                                st.json(result)
                            else:
                                st.error(f"Failed: {result}")
                        except Exception as e:
                            st.error(f"Error: {str(e)}")
            
            # 14. Cognitive Assessment
            with st.expander("1️⃣4️⃣ Cognitive Assessment", expanded=False):
                json14_data = extracted_data.get('cognitive_assessment', {}).copy()
                if st.session_state.special_education_id:
                    json14_data['special_education'] = st.session_state.special_education_id
                st.json(json14_data)
                
                edited_json14 = st.text_area(
                    "Edit Cognitive Assessment:",
                    value=json.dumps(json14_data, indent=2, ensure_ascii=False),
                    height=300,
                    key="json14_editor"
                )
                
                if st.button("Push Cognitive Assessment", type="primary", key="push_cognitive"):
                    if not st.session_state.special_education_id:
                        st.error("⚠️ Please push Special Education Assessment first!")
                    else:
                        try:
                            json14_parsed = json.loads(edited_json14)
                            json14_parsed['special_education'] = st.session_state.special_education_id
                            
                            success, result = make_api_request(
                                API_ENDPOINTS["cognitive_assessment"],
                                json14_parsed,
                                bearer_token
                            )
                            
                            if success:
                                st.success("✅ Cognitive Assessment pushed!")
                                st.json(result)
                            else:
                                st.error(f"Failed: {result}")
                        except Exception as e:
                            st.error(f"Error: {str(e)}")
            
            # 15. Other Assessments
            with st.expander("1️⃣5️⃣ Other Assessments", expanded=False):
                json15_data = extracted_data.get('other_assessments', {}).copy()
                if st.session_state.special_education_id:
                    json15_data['special_education'] = st.session_state.special_education_id
                st.json(json15_data)
                
                edited_json15 = st.text_area(
                    "Edit Other Assessments:",
                    value=json.dumps(json15_data, indent=2, ensure_ascii=False),
                    height=300,
                    key="json15_editor"
                )
                
                if st.button("Push Other Assessments", type="primary", key="push_other_assessments"):
                    if not st.session_state.special_education_id:
                        st.error("⚠️ Please push Special Education Assessment first!")
                    else:
                        try:
                            json15_parsed = json.loads(edited_json15)
                            json15_parsed['special_education'] = st.session_state.special_education_id
                            
                            success, result = make_api_request(
                                API_ENDPOINTS["other_assessments"],
                                json15_parsed,
                                bearer_token
                            )
                            
                            if success:
                                st.success("✅ Other Assessments pushed!")
                                st.json(result)
                            else:
                                st.error(f"Failed: {result}")
                        except Exception as e:
                            st.error(f"Error: {str(e)}")
            
            st.divider()
            
            st.subheader("📊 Assessment Dependent Fields (16-17)")
            st.info("⚠️ These fields require Assessment ID")
            
            # 16. Basic Observation
            with st.expander("1️⃣6️⃣ Basic Observation", expanded=False):
                json16_data = extracted_data.get('basic_observation', {}).copy()
                if st.session_state.assessment_id:
                    json16_data['assessment'] = st.session_state.assessment_id
                st.json(json16_data)
                
                edited_json16 = st.text_area(
                    "Edit Basic Observation:",
                    value=json.dumps(json16_data, indent=2, ensure_ascii=False),
                    height=300,
                    key="json16_editor"
                )
                
                if st.button("Push Basic Observation", type="primary", key="push_basic_observation"):
                    if not st.session_state.assessment_id:
                        st.error("⚠️ Please push Assessment first!")
                    else:
                        try:
                            json16_parsed = json.loads(edited_json16)
                            json16_parsed['assessment'] = st.session_state.assessment_id
                            
                            success, result = make_api_request(
                                API_ENDPOINTS["basic_observation"],
                                json16_parsed,
                                bearer_token
                            )
                            
                            if success:
                                st.success("✅ Basic Observation pushed!")
                                st.json(result)
                            else:
                                st.error(f"Failed: {result}")
                        except Exception as e:
                            st.error(f"Error: {str(e)}")
            
            # 17. Sensory Assessment
            with st.expander("1️⃣7️⃣ Sensory Assessment", expanded=False):
                json17_data = extracted_data.get('sensory_assessment', {}).copy()
                if st.session_state.assessment_id:
                    json17_data['assessment'] = st.session_state.assessment_id
                st.json(json17_data)
                
                edited_json17 = st.text_area(
                    "Edit Sensory Assessment:",
                    value=json.dumps(json17_data, indent=2, ensure_ascii=False),
                    height=300,
                    key="json17_editor"
                )
                
                if st.button("Push Sensory Assessment", type="primary", key="push_sensory_assessment"):
                    if not st.session_state.assessment_id:
                        st.error("⚠️ Please push Assessment first!")
                    else:
                        try:
                            json17_parsed = json.loads(edited_json17)
                            json17_parsed['assessment'] = st.session_state.assessment_id
                            
                            success, result = make_api_request(
                                API_ENDPOINTS["sensory_assessment"],
                                json17_parsed,
                                bearer_token
                            )
                            
                            if success:
                                st.success("✅ Sensory Assessment pushed!")
                                st.json(result)
                            else:
                                st.error(f"Failed: {result}")
                        except Exception as e:
                            st.error(f"Error: {str(e)}")
            
            st.divider()
            
            st.subheader("🆔 Generated IDs")
            id_col1, id_col2, id_col3, id_col4, id_col5 = st.columns(5)
            
            with id_col1:
                if st.session_state.student_id:
                    st.success(f"**Student ID:**\n{st.session_state.student_id}")
                else:
                    st.info("**Student ID:**\nNot generated")
            
            with id_col2:
                if st.session_state.assessment_id:
                    st.success(f"**Assessment ID:**\n{st.session_state.assessment_id}")
                else:
                    st.info("**Assessment ID:**\nNot generated")
            
            with id_col3:
                if st.session_state.pre_assessment_id:
                    st.success(f"**Pre-Assessment ID:**\n{st.session_state.pre_assessment_id}")
                else:
                    st.info("**Pre-Assessment ID:**\nNot generated")
            
            with id_col4:
                if st.session_state.speech_assessment_id:
                    st.success(f"**Speech Assessment ID:**\n{st.session_state.speech_assessment_id}")
                else:
                    st.info("**Speech Assessment ID:**\nNot generated")
            
            with id_col5:
                if st.session_state.special_education_id:
                    st.success(f"**Special Education ID:**\n{st.session_state.special_education_id}")
                else:
                    st.info("**Special Education ID:**\nNot generated")
            
            st.divider()
            
            st.subheader("📥 Download Extracted Data")
            st.download_button(
                label="📥 Download All Extracted Data (JSON)",
                data=json.dumps(extracted_data, indent=2, ensure_ascii=False),
                file_name="assessment_all_extracted_data.json",
                mime="application/json"
            )
        
        except Exception as e:
            st.error(f"❌ Error processing file: {str(e)}")
            st.exception(e)


if __name__ == "__main__":
    main()
